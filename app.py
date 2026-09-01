import os
import re
import json
import logging
import time
from collections import defaultdict
from io import BytesIO
from html import escape
from datetime import datetime
from flask import Flask, request, redirect, render_template_string, session, Response, send_file, send_from_directory
from flask_sqlalchemy import SQLAlchemy
from flask_wtf.csrf import CSRFProtect
from werkzeug.security import generate_password_hash, check_password_hash

# ReportLab, Docx, & Ebooklib untuk Fitur Ekspor Dokumen
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from docx import Document
from ebooklib import epub

app = Flask(__name__)
app.secret_key = 'karya-dede-suhendra-secret-key-2026-upgraded'

# ==================================================
# KONFIGURASI KEAMANAN CSRF
# ==================================================
csrf = CSRFProtect(app)

login_attempts = defaultdict(list)

def cek_rate_limit_login(ip):
    sekarang = time.time()
    login_attempts[ip] = [t for t in login_attempts[ip] if sekarang - t < 60]
    if len(login_attempts[ip]) >= 5:
        return False
    login_attempts[ip].append(sekarang)
    return True

# Helper Pembersih Teks PDF (Prioritas Tinggi: Anti-Crash PDF)
def bersihkan_teks_pdf(teks):
    if not teks:
        return ""
    teks_bersih = re.sub(r'<[^>]+>', '', teks)
    return escape(teks_bersih).replace('\n', '<br/>')

# ==================================================
# KONFIGURASI LOGGING SISTEM
# ==================================================
logging.basicConfig(level=logging.INFO, format='%(asctime)s [%(levelname)s] %(message)s')

# ==================================================
# KONFIGURASI FOLDER PENYIMPANAN PDF AMAN
# ==================================================
UPLOAD_PDF_FOLDER = os.path.join(os.path.abspath(os.path.dirname(__file__)), 'uploaded_pdfs')
os.makedirs(UPLOAD_PDF_FOLDER, exist_ok=True)

# ==================================================
# KONFIGURASI DATABASE
# ==================================================
TURSO_DATABASE_URL = os.environ.get('TURSO_DATABASE_URL')
TURSO_AUTH_TOKEN = os.environ.get('TURSO_AUTH_TOKEN')

if TURSO_DATABASE_URL and TURSO_AUTH_TOKEN:
    db_path = os.path.join(os.path.abspath(os.path.dirname(__file__)), 'karya_buku.db')
    app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{db_path}'
    database_info = "SQLite Railway (Stabil)"
else:
    db_path = os.path.join(os.path.abspath(os.path.dirname(__file__)), 'karya_buku.db')
    app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{db_path}'
    database_info = "SQLite Lokal"

app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['PERMANENT_SESSION_LIFETIME'] = 86400

db = SQLAlchemy(app)

ADMIN_USER = "dede"
ADMIN_PASS_HASH = generate_password_hash("suhendra123")

# ==================================================
# MODEL DATABASE
# ==================================================
class Tema(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    nama = db.Column(db.String(100), nullable=False)
    dibuat_pada = db.Column(db.DateTime, default=datetime.utcnow)
    diupdate_pada = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    buku_list = db.relationship('Buku', backref='tema', lazy=True, cascade="all, delete-orphan")

class Buku(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    judul = db.Column(db.String(200), nullable=False)
    subjudul = db.Column(db.String(250), nullable=True)
    kutipan = db.Column(db.Text, nullable=True)
    tema_id = db.Column(db.Integer, db.ForeignKey('tema.id'), nullable=False)
    status = db.Column(db.String(20), default='selesai')
    dibuat_pada = db.Column(db.DateTime, default=datetime.utcnow)
    diupdate_pada = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    catatan_list = db.relationship('Catatan', backref='buku', lazy=True, cascade="all, delete-orphan")

class Catatan(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    bagian = db.Column(db.String(100), nullable=True)
    judul_bab = db.Column(db.String(200), nullable=False)
    isi = db.Column(db.Text, nullable=False)
    urutan = db.Column(db.Integer, default=1)
    buku_id = db.Column(db.Integer, db.ForeignKey('buku.id'), nullable=False)
    dibuat_pada = db.Column(db.DateTime, default=datetime.utcnow)
    diupdate_pada = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    favorit = db.Column(db.Boolean, default=False)
    tag = db.Column(db.String(200), nullable=True)

class EsaiPenulis(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    judul = db.Column(db.String(200), nullable=False)
    kategori = db.Column(db.String(100), default="Refleksi Harian")
    isi = db.Column(db.Text, nullable=False)
    tanggal = db.Column(db.String(50), nullable=True)
    dibuat_pada = db.Column(db.DateTime, default=datetime.utcnow)
    diupdate_pada = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    favorit = db.Column(db.Boolean, default=False)

class TongSampah(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    tipe = db.Column(db.String(20), nullable=False)
    data_json = db.Column(db.Text, nullable=False)
    dihapus_pada = db.Column(db.DateTime, default=datetime.utcnow)

class LogAktivitas(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    aksi = db.Column(db.String(200), nullable=False)
    keterangan = db.Column(db.Text, nullable=True)
    waktu = db.Column(db.DateTime, default=datetime.utcnow)

class PenandaBaca(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    buku_id = db.Column(db.Integer, db.ForeignKey('buku.id'), nullable=False)
    catatan_id = db.Column(db.Integer, db.ForeignKey('catatan.id'), nullable=False)
    waktu_baca = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    buku = db.relationship('Buku', backref=db.backref('penanda', uselist=False, cascade="all, delete-orphan"))
    catatan = db.relationship('Catatan')

with app.app_context():
    db.create_all()
    if Tema.query.count() == 0:
        db.session.add_all([
            Tema(nama='Filsafat'),
            Tema(nama='Keuangan'),
            Tema(nama='Komunikasi')
        ])
        db.session.commit()
    
    if Buku.query.count() == 0 and os.path.exists('backup_karya.json'):
        try:
            with open('backup_karya.json', 'r', encoding='utf-8') as f:
                data = json.load(f)
                for t in data.get('tema', []):
                    if not Tema.query.get(t['id']):
                        db.session.add(Tema(id=t['id'], nama=t['nama']))
                for b in data.get('buku', []):
                    if not Buku.query.get(b['id']):
                        db.session.add(Buku(id=b['id'], judul=b['judul'], subjudul=b.get('subjudul'), tema_id=b['tema_id'], kutipan=b.get('kutipan')))
                for c in data.get('catatan', []):
                    if not Catatan.query.get(c['id']):
                        db.session.add(Catatan(id=c['id'], bagian=c.get('bagian'), judul_bab=c['judul_bab'], isi=c['isi'], buku_id=c['buku_id'], urutan=c.get('urutan', 1)))
                for e in data.get('esai', []):
                    if not EsaiPenulis.query.get(e['id']):
                        db.session.add(EsaiPenulis(id=e['id'], judul=e['judul'], kategori=e.get('kategori', 'Refleksi'), isi=e['isi']))
                db.session.commit()
                app.logger.info("Auto-restore database dari file backup_karya.json berhasil dilakukan.")
        except Exception as err:
            db.session.rollback()
            app.logger.error(f"Gagal melakukan auto-restore: {str(err)}")

    app.logger.info(f"Database berhasil diinisialisasi menggunakan: {database_info}")

# ==================================================
# ROUTE FILE MEDIA & STATIC (LOGO & MANIFEST PWA)
# ==================================================
@app.route('/profile.jpg')
def serve_profile():
    return send_from_directory('.', 'profile.jpg')

@app.route('/logo.png')
def serve_logo():
    return send_from_directory('static', 'logo.png')

@app.route('/manifest.json')
def serve_manifest():
    return send_from_directory('static', 'manifest.json')

def masukkan_sampah(tipe, data_dict):
    item = TongSampah(tipe=tipe, data_json=json.dumps(data_dict, ensure_ascii=False))
    db.session.add(item)
    db.session.commit()

def catat_log(aksi, keterangan=""):
    log_baru = LogAktivitas(aksi=aksi, keterangan=keterangan)
    db.session.add(log_baru)
    db.session.commit()

@app.errorhandler(404)
def halaman_tidak_ditemukan(e):
    return """
    <!DOCTYPE html>
    <html lang="id">
    <head><meta charset="UTF-8"><title>404</title><link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet"></head>
    <body style="background:#0b132b; color:#f8fafc; display:flex; justify-content:center; align-items:center; min-height:100vh; text-align:center;">
        <div style="max-width:450px; padding:30px; background:#1c2541; border:1px solid #334155; border-radius:18px;">
            <h1 class="text-warning fw-bold mb-2" style="font-size: 64px;">404</h1>
            <h4 class="mb-3 fw-bold">Halaman Tidak Ditemukan</h4>
            <a href="/" class="btn btn-outline-warning rounded-pill px-4 fw-bold text-white">&larr; Kembali ke Beranda</a>
        </div>
    </body>
    </html>
    """, 404

@app.errorhandler(500)
def kesalahan_server(e):
    return f"""
    <!DOCTYPE html>
    <html lang="id">
    <head><meta charset="UTF-8"><title>500</title><link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet"></head>
    <body style="background:#0b132b; color:#f8fafc; display:flex; justify-content:center; align-items:center; min-height:100vh; text-align:center;">
        <div style="max-width:500px; padding:30px; background:#1c2541; border:1px solid #334155; border-radius:18px;">
            <h1 class="text-danger fw-bold mb-2" style="font-size: 64px;">500</h1>
            <h4 class="mb-3 fw-bold">Terjadi Kesalahan Server</h4>
            <div class="bg-dark p-2 rounded text-start small text-warning mb-4" style="max-height:120px; overflow-y:auto;">Detail: {escape(str(e))}</div>
            <a href="/" class="btn btn-outline-light rounded-pill px-4 fw-bold">&larr; Kembali ke Beranda</a>
        </div>
    </body>
    </html>
    """, 500

CSS_SHARED = """
    :root { 
        --bg-paper: #f7f4ef; --card-paper: #ffffff; --text-main: #1e293b; --text-muted: #64748b;
        --gold-gradient: linear-gradient(135deg, #bf953f, #fcf6ba, #b38728, #fbf5b7);
        --border-color: rgba(212, 175, 55, 0.4); --input-bg: #ffffff; --input-text: #1e293b;
    }
    [data-theme="gelap"] {
        --bg-paper: #0b132b; --card-paper: #1c2541; --text-main: #f8fafc; --text-muted: #cbd5e1;
        --border-color: #334155; --input-bg: #1e293b; --input-text: #f8fafc;
    }
    body { background-color: var(--bg-paper) !important; font-family: 'Plus Jakarta Sans', sans-serif; color: var(--text-main) !important; transition: background-color 0.2s ease, color 0.2s ease; }
    h1, h2, h3, h4, h5, h6, p, div, label, span, small { color: inherit !important; }
    .text-muted { color: var(--text-muted) !important; }
    .header-title { font-family: 'Cinzel', serif; font-weight: 700; color: var(--text-main) !important; }
    .top-badge { background: #fcf8ec; color: #b38728 !important; border: 1px solid rgba(212,175,55,0.4); font-weight:700; padding:6px 16px; border-radius:30px; font-size:0.8rem; }
    .card-gold { background:var(--card-paper) !important; border-radius:18px; border:1px solid var(--border-color) !important; box-shadow:0 10px 25px rgba(0,0,0,0.03); transition:all 0.3s ease; position:relative; overflow:hidden; }
    .card-gold::before { content:''; position:absolute; top:0; left:0; width:100%; height:4px; background:var(--gold-gradient); }
    .card-gold:hover { transform:translateY(-4px); box-shadow:0 15px 30px rgba(212,175,55,0.18); }
    .btn-custom-outline { color: var(--text-main) !important; border-color: var(--border-color) !important; background: var(--card-paper) !important; }
    .btn-custom-outline:hover { background: #b38728 !important; color: #fff !important; }
    .search-box { position:relative; }
    .search-box input { padding-left:38px; background: var(--input-bg) !important; color: var(--input-text) !important; border-color: var(--border-color) !important; }
    .search-box i { position:absolute; left:12px; top:50%; transform:translateY(-50%); color:#888; }
    .badge-count { background:#b38728; color:white !important; border-radius:50%; width:22px; height:22px; display:inline-flex; align-items:center; justify-content:center; font-size:11px; margin-left:6px; }
    .btn-mode-toggle { position:fixed; top:15px; right:15px; z-index:100; border-radius:50%; width:44px; height:44px; display:flex; align-items:center; justify-content:center; background: var(--card-paper) !important; border: 2px solid #b38728 !important; color: #b38728 !important; box-shadow: 0 4px 10px rgba(0,0,0,0.2); }
    
    .btn-to-top {
        position: fixed; bottom: 25px; right: 25px; z-index: 99;
        background: #b38728; color: white; border: none; border-radius: 50%;
        width: 48px; height: 48px; display: flex; align-items: center; justify-content: center;
        box-shadow: 0 4px 12px rgba(0,0,0,0.3); opacity: 0; transition: opacity 0.3s, transform 0.2s;
        cursor: pointer; pointer-events: none; text-decoration: none;
    }
    .btn-to-top.show { opacity: 1; pointer-events: auto; }
    .btn-to-top:hover { transform: scale(1.1); background: #96701f; color: white; }

    #reading-progress {
        position: fixed; top: 0; left: 0; height: 4px;
        background: linear-gradient(90deg, #bf953f, #fcf6ba, #b38728);
        width: 0%; z-index: 9999; transition: width 0.1s ease-out;
    }

    .note-card-badge { display: block; width: fit-content; background: rgba(56, 189, 248, 0.15); color: #38bdf8 !important; font-size: 11px; font-weight: 700; padding: 5px 12px; border-radius: 6px; text-transform: uppercase; margin-bottom: 8px; }
    .note-card-title { color: #b38728 !important; font-size: 20px; font-weight: 800; margin-top: 4px; margin-bottom: 12px; padding-right: 170px; }
    .markdown-body img {
        max-width: 100% !important;
        height: auto !important;
        border-radius: 12px;
        margin: 15px auto;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.15);
        display: block;
    
    .toc-box { background: rgba(179, 135, 40, 0.05); border: 1px dashed var(--border-color); border-radius: 12px; padding: 20px; margin-bottom: 25px; }
    .toc-section-title { font-size: 13px; font-weight: 800; text-transform: uppercase; color: #b38728; letter-spacing: 0.5px; margin-top: 14px; margin-bottom: 6px; border-bottom: 1px solid rgba(179, 135, 40, 0.2); padding-bottom: 3px; }
    .toc-section-title:first-child { margin-top: 0; }
    .toc-list { list-style-type: none; padding-left: 0; margin-bottom: 10px; }
    .toc-list li { margin-bottom: 5px; font-size: 13.5px; padding-left: 12px; }
    .toc-list a { color: var(--text-main); text-decoration: none; font-weight: 600; display: inline-flex; align-items: center; gap: 6px; }
    .toc-list a:hover { color: #b38728; text-decoration: underline; }

    .section-header-card { background: rgba(179, 135, 40, 0.12); border-left: 5px solid #b38728; border-radius: 10px; padding: 12px 18px; margin-top: 30px; margin-bottom: 15px; font-family: 'Cinzel', serif; font-weight: 800; color: #b38728 !important; font-size: 16px; text-transform: uppercase; }

    /* SPLASH SCREEN & NATIVE HEADER APP */
    #splash-screen {
        position: fixed; top: 0; left: 0; width: 100%; height: 100vh;
        background: radial-gradient(circle at center, #1c2541 0%, #0b132b 100%);
        display: flex; flex-direction: column; justify-content: center; align-items: center;
        z-index: 999999; transition: opacity 0.8s cubic-bezier(0.4, 0, 0.2, 1), visibility 0.8s;
    }
    #splash-screen.hidden { opacity: 0; visibility: hidden; pointer-events: none; }
    .splash-container { text-align: center; animation: zoomInSplash 0.8s ease-out; }
    .splash-logo-wrapper { position: relative; width: 130px; height: 130px; margin: 0 auto 20px auto; }
    .splash-logo-img {
        width: 100%; height: 100%; border-radius: 50%; object-fit: cover;
        border: 3px solid #b38728; box-shadow: 0 0 25px rgba(212, 175, 55, 0.5);
        animation: logoGlow 2s infinite alternate ease-in-out;
    }
    .splash-title { font-family: 'Cinzel', serif; font-weight: 800; font-size: 22px; color: #fcf6ba !important; letter-spacing: 2px; margin-bottom: 5px; }
    .splash-subtitle { font-size: 12px; color: #94a3b8 !important; letter-spacing: 3px; text-transform: uppercase; font-weight: 600; }
    @keyframes logoGlow {
        0% { transform: scale(0.96); box-shadow: 0 0 15px rgba(212, 175, 55, 0.3); }
        100% { transform: scale(1.04); box-shadow: 0 0 35px rgba(212, 175, 55, 0.7); }
    }
    @keyframes zoomInSplash { from { opacity: 0; transform: scale(0.85); } to { opacity: 1; transform: scale(1); } }
    .app-header-logo {
        width: 100px; height: 100px; border-radius: 50%; object-fit: cover;
        border: 3px solid #b38728; padding: 3px; background: var(--card-paper);
        box-shadow: 0 8px 20px rgba(212, 175, 55, 0.25); transition: transform 0.3s ease;
    }
    .app-header-logo:hover { transform: rotate(5deg) scale(1.05); }
"""

JS_THEME_SCRIPT = """
<script>
    (function() {
        const savedTheme = localStorage.getItem('theme_mode') || 'terang';
        document.documentElement.setAttribute('data-theme', savedTheme);
    })();

    document.addEventListener("DOMContentLoaded", function() {
        updateIcon();
        setupScrollTopBtn();
        window.addEventListener('scroll', function() {
            const winScroll = document.body.scrollTop || document.documentElement.scrollTop;
            const height = document.documentElement.scrollHeight - document.documentElement.clientHeight;
            const scrolled = (winScroll / height) * 100;
            const progressBar = document.getElementById("reading-progress");
            if (progressBar) {
                progressBar.style.width = scrolled + "%";
            }
        });
    });

    window.addEventListener('load', function() {
        const splash = document.getElementById('splash-screen');
        if (splash) {
            setTimeout(function() {
                splash.classList.add('hidden');
            }, 1200);
        }
    });

    function toggleModeInstan() {
        const currentTheme = document.documentElement.getAttribute('data-theme') === 'gelap' ? 'terang' : 'gelap';
        document.documentElement.setAttribute('data-theme', currentTheme);
        localStorage.setItem('theme_mode', currentTheme);
        updateIcon();
    }

    function updateIcon() {
        const icon = document.getElementById('icon-mode');
        if (icon) {
            const currentTheme = document.documentElement.getAttribute('data-theme');
            icon.className = currentTheme === 'gelap' ? 'fa-solid fa-sun text-warning' : 'fa-solid fa-moon text-dark';
        }
    }

    function setupScrollTopBtn() {
        const btn = document.getElementById('btnScrollTop');
        if (!btn) return;
        window.addEventListener('scroll', function() {
            if (window.pageYOffset > 300) {
                btn.classList.add('show');
            } else {
                btn.classList.remove('show');
            }
        });
    }

    let deferredPrompt;
    window.addEventListener('beforeinstallprompt', (e) => {
        e.preventDefault();
        deferredPrompt = e;
    });

    function manualInstallGuide() {
        if (deferredPrompt) {
            deferredPrompt.prompt();
        } else {
            alert("Untuk menginstal aplikasi PWA di HP Anda:\\n\\n1. Ketuk titik tiga (⋮) di pojok kanan atas Chrome.\\n2. Pilih 'Tambahkan ke Layar Utama' / 'Instal Aplikasi'.");
        }
    }
</script>
"""

HTML_INDEX = """
<!DOCTYPE html>
<html lang="id">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Ruang Literasi Hendra - Dede Suhendra</title>
    <link rel="manifest" href="/static/manifest.json">
    <meta name="theme-color" content="#b38728">
    <meta name="mobile-web-app-capable" content="yes">
    <meta name="apple-mobile-web-app-capable" content="yes">
    <meta name="apple-mobile-web-app-status-bar-style" content="black-translucent">
    <link rel="apple-touch-icon" href="/static/logo.png">
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
    <link href="https://fonts.googleapis.com/css2?family=Cinzel:wght@600;700;800&family=Plus+Jakarta+Sans:wght@400;600;700;800&display=swap" rel="stylesheet">
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css">
    <style>""" + CSS_SHARED + """</style>
    """ + JS_THEME_SCRIPT + """
</head>
<body>
<div id="splash-screen">
    <div class="splash-container">
        <div class="splash-logo-wrapper">
            <img src="/static/logo.png" class="splash-logo-img" alt="Ruang Literasi Hendra">
        </div>
        <div class="splash-title">RUANG LITERASI HENDRA</div>
        <div class="splash-subtitle mb-3">Koleksi Karya & Refleksi</div>
        <div class="spinner-border text-warning spinner-border-sm" role="status"></div>
    </div>
</div>

<button class="btn btn-mode-toggle" onclick="toggleModeInstan()" title="Ganti Mode Tampilan">
    <i class="fa-solid fa-moon" id="icon-mode"></i>
</button>

<a href="#" class="btn-to-top" id="btnScrollTop" title="Kembali ke Atas"><i class="fa-solid fa-arrow-up"></i></a>

<div class="container py-5" style="max-width:850px;">
    <div class="d-flex justify-content-between align-items-center mb-4 flex-wrap gap-2">
        <a href="/penulis" class="btn btn-custom-outline rounded-pill btn-sm px-3 fw-bold"><i class="fa-solid fa-user-tie me-1"></i> Profil Penulis</a>
        <div class="d-flex gap-2 align-items-center flex-wrap">
            <a href="/favorit" class="btn btn-custom-outline rounded-pill btn-sm fw-bold"><i class="fa-solid fa-star me-1 text-warning"></i> Favorit</a>
            <a href="/baca-pdf" class="btn btn-custom-outline rounded-pill btn-sm fw-bold"><i class="fa-solid fa-book-open me-1 text-danger"></i> PDF Reader</a>
            <a href="/statistik" class="btn btn-custom-outline rounded-pill btn-sm fw-bold"><i class="fa-solid fa-chart-simple me-1 text-info"></i> Statistik</a>
            <a href="/tong-sampah" class="btn btn-custom-outline rounded-pill btn-sm fw-bold"><i class="fa-solid fa-trash-can me-1 text-secondary"></i> Tong Sampah</a>
            {% if is_admin %}
                <a href="/riwayat-log" class="btn btn-custom-outline rounded-pill btn-sm fw-bold"><i class="fa-solid fa-list-check me-1 text-warning"></i> Log</a>
                <a href="/backup" class="btn btn-custom-outline rounded-pill btn-sm fw-bold"><i class="fa-solid fa-download me-1 text-primary"></i> Backup</a>
                <button type="button" class="btn btn-custom-outline rounded-pill btn-sm fw-bold text-success" data-bs-toggle="modal" data-bs-target="#restoreModal"><i class="fa-solid fa-upload me-1"></i> Restore</button>
                <a href="/logout" class="btn btn-warning rounded-pill btn-sm fw-bold text-dark"><i class="fa-solid fa-right-from-bracket me-1"></i> Logout</a>
            {% else %}
                <a href="/login" class="btn btn-custom-outline rounded-pill btn-sm fw-bold"><i class="fa-solid fa-lock me-1"></i> Login</a>
            {% endif %}
        </div>
    </div>

    {% if penanda_list %}
    <div class="card-gold p-3 mb-4 rounded-4 border-warning">
        <div class="d-flex align-items-center justify-content-between flex-wrap gap-2">
            <div class="d-flex align-items-center">
                <div class="bg-warning text-dark rounded-circle p-2 me-3"><i class="fa-solid fa-bookmark fs-5"></i></div>
                <div>
                    <span class="badge bg-warning text-dark fw-bold mb-1">LANJUTKAN MEMBACA</span>
                    <h6 class="mb-0 fw-bold">Posisi Terakhir: {{ penanda_list[0].catatan.judul_bab }}</h6>
                    <small class="text-muted">Buku: <strong>{{ penanda_list[0].buku.judul }}</strong></small>
                </div>
            </div>
            <a href="/buku/{{ penanda_list[0].buku_id }}" class="btn btn-warning btn-sm text-dark rounded-pill fw-bold px-3">
                <i class="fa-solid fa-play me-1"></i> Lanjut Baca
            </a>
        </div>
    </div>
    {% endif %}

    <div class="modal fade" id="restoreModal" tabindex="-1" aria-hidden="true">
      <div class="modal-dialog">
        <div class="modal-content bg-dark text-white border-secondary">
          <div class="modal-header border-secondary">
            <h5 class="modal-title fw-bold">📥 Restore Data Catatan</h5>
            <button type="button" class="btn-close btn-close-white" data-bs-dismiss="modal"></button>
          </div>
          <form action="/restore" method="POST" enctype="multipart/form-data">
              <input type="hidden" name="csrf_token" value="{{ csrf_token() }}">
              <div class="modal-body">
                <p class="small text-muted">Upload file <code>backup_karya.json</code> untuk mengembalikan semua naskah.</p>
                <input type="file" name="file_json" class="form-control bg-secondary text-white" accept=".json" required>
              </div>
              <div class="modal-footer border-secondary">
                <button type="button" class="btn btn-sm btn-outline-light" data-bs-dismiss="modal">Batal</button>
                <button type="submit" class="btn btn-sm btn-success fw-bold">Upload & Restore</button>
              </div>
          </form>
        </div>
      </div>
    </div>
    
    <div class="text-center mb-4">
        <img src="/static/logo.png" alt="Ruang Literasi Hendra" class="app-header-logo mb-3">
        <h1 class="header-title h2 mb-1" style="font-family:'Cinzel', serif;">RUANG LITERASI HENDRA</h1>
        <span class="top-badge mb-2 d-inline-block">OFFICIAL DIGITAL VAULT</span>
        <p class="text-muted small mt-1">Disusun & Dikelola oleh <strong>Dede Suhendra</strong></p>
    </div>

    <div id="pwa-install-container" class="text-center mb-4">
        <button id="btnInstallPWA" class="btn btn-warning text-dark rounded-pill px-4 py-2 fw-bold shadow-sm" onclick="manualInstallGuide()">
            <i class="fa-solid fa-download me-2"></i> Download / Instal Aplikasi
        </button>
    </div>

    <div class="card-gold p-3 mb-4 rounded-3">
        <form action="/cari" method="GET" class="search-box mb-2">
            <i class="fa-solid fa-magnifying-glass"></i>
            <input type="text" name="q" class="form-control form-control-sm" placeholder="Cari kata kunci dalam isi bab/naskah (Tekan Enter)...">
        </form>
        <div class="d-flex flex-wrap gap-2 mt-2">
            <span class="badge border filter-tag active" data-filter="semua" onclick="filterKategori('semua')">Semua</span>
            {% for t in tema_list %}
            <span class="badge border filter-tag" data-filter="{{ t.nama }}" onclick="filterKategori('{{ t.nama }}')">{{ t.nama }}</span>
            {% endfor %}
        </div>
    </div>

    {% if is_admin %}
    <div class="card-gold p-3 mb-4 rounded-3">
        <h6 class="fw-bold text-success mb-2"><i class="fa-solid fa-folder-plus me-1"></i> Tambah Tema Baru</h6>
        <form action="/tambah-tema" method="POST" class="row g-2">
            <input type="hidden" name="csrf_token" value="{{ csrf_token() }}">
            <div class="col-9"><input type="text" name="nama" class="form-control form-control-sm" placeholder="Misal: Psikologi..." required></div>
            <div class="col-3"><button type="submit" class="btn btn-success btn-sm w-100 fw-bold">Tambah</button></div>
        </form>
    </div>
    {% endif %}

    <h5 class="fw-bold mb-3" style="font-family:'Cinzel',serif;">Pilih Kategori Karya:</h5>
    <div class="row g-3" id="daftar-kategori">
        <div class="col-12 kategori-item" data-nama="Jurnal & Esai" data-keywords="jurnal esai refleksi harian">
            <a href="/catatan-penulis" class="text-decoration-none">
                <div class="card-gold p-4">
                    <div class="d-flex justify-content-between align-items-center">
                        <div class="d-flex align-items-center">
                            <div class="bg-warning text-dark rounded-circle p-3 me-3"><i class="fa-solid fa-feather-pointed fs-4"></i></div>
                            <div>
                                <span class="badge bg-warning text-dark fw-bold mb-1">Ruang Refleksi</span>
                                <h4 class="h5 mb-1 fw-bold">✍️ Jurnal & Esai Bebas <span class="badge-count">{{ jumlah_esai }}</span></h4>
                                <p class="text-muted small mb-0">Kumpulan perenungan harian & ide acak.</p>
                            </div>
                        </div>
                        <i class="fa-solid fa-chevron-right text-warning fs-5"></i>
                    </div>
                </div>
            </a>
        </div>

        {% for tema in tema_list %}
        <div class="col-md-6 kategori-item" data-nama="{{ tema.nama }}" data-keywords="{{ tema.nama | lower }}">
            <div class="card-gold p-4">
                <div class="d-flex justify-content-between align-items-center mb-2">
                    <a href="/tema/{{ tema.id }}" class="text-decoration-none d-flex align-items-center flex-grow-1">
                        <i class="fa-solid fa-folder-open text-warning me-3 fs-4"></i>
                        <h4 class="h5 mb-0 fw-bold">{{ tema.nama }} <span class="badge-count">{{ tema.buku_list|length }}</span></h4>
                    </a>
                    {% if is_admin %}
                    <form action="/hapus-tema/{{ tema.id }}" method="POST" class="mb-0" onsubmit="return confirm('Hapus beserta semua isinya?');">
                        <input type="hidden" name="csrf_token" value="{{ csrf_token() }}">
                        <button type="submit" class="btn btn-link text-danger p-0 ms-2"><i class="fa-solid fa-trash"></i></button>
                    </form>
                    {% endif %}
                </div>
                <small class="text-muted ms-4">Dibuat: {{ tema.dibuat_pada.strftime('%d %b %Y') if tema.dibuat_pada else '-' }} &rarr;</small>
            </div>
        </div>
        {% endfor %}
    </div>
</div>

<script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/js/bootstrap.bundle.min.js"></script>
<script>
function jalankanPencarianCanggih(kata) {
    const keyword = kata.toLowerCase().trim();
    document.querySelectorAll('.kategori-item').forEach(item => {
        const keywords = item.getAttribute('data-keywords') || '';
        const textContent = item.textContent.toLowerCase();
        if (keyword === '' || keywords.includes(keyword) || textContent.includes(keyword)) {
            item.style.display = '';
        } else {
            item.style.display = 'none';
        }
    });
}
function filterKategori(nama) {
    document.querySelectorAll('.filter-tag').forEach(t => t.classList.remove('active'));
    document.querySelector(`[data-filter="${nama}"]`).classList.add('active');
    document.querySelectorAll('.kategori-item').forEach(item => {
        item.style.display = nama === 'semua' || item.getAttribute('data-nama') === nama ? '' : 'none';
    });
}
</script>
</body>
</html>
"""

HTML_HASIL_CARI = """
<!DOCTYPE html>
<html lang="id">
<head>
    <meta charset="UTF-8">
    <title>Hasil Pencarian: {{ query }}</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css">
    <style>""" + CSS_SHARED + """</style>
    """ + JS_THEME_SCRIPT + """
</head>
<body>
<button class="btn btn-mode-toggle" onclick="toggleModeInstan()"><i class="fa-solid fa-moon" id="icon-mode"></i></button>
<div class="container py-5" style="max-width:800px;">
    <a href="/" class="btn btn-custom-outline rounded-pill btn-sm mb-4 fw-bold">&larr; Kembali</a>
    
    <div class="card-gold p-4 mb-4 rounded-4">
        <h4 class="fw-bold mb-2">🔍 Hasil Pencarian: "<span class="text-warning">{{ query }}</span>"</h4>
        <p class="text-muted small mb-0">Ditemukan {{ hasil_catatan|length }} bab buku dan {{ hasil_esai|length }} esai.</p>
    </div>

    <h5 class="fw-bold text-warning mb-3">Bab Buku</h5>
    {% for c in hasil_catatan %}
    <div class="card-gold p-3 mb-3">
        <div class="d-flex justify-content-between align-items-center mb-1">
            <span class="badge bg-secondary">{{ c.bagian or 'Bab' }}</span>
            <a href="/buku/{{ c.buku_id }}#bab-{{ c.id }}" class="btn btn-warning btn-sm text-dark fw-bold rounded-pill">Buka Bab &rarr;</a>
        </div>
        <h5 class="fw-bold mb-1">{{ c.judul_bab }}</h5>
        <p class="small text-muted mb-0">{{ c.isi[:180] }}...</p>
    </div>
    {% else %}
    <p class="text-muted small mb-4">Tidak ada bab yang cocok.</p>
    {% endfor %}

    <h5 class="fw-bold text-warning mb-3">Jurnal & Esai</h5>
    {% for e in hasil_esai %}
    <div class="card-gold p-3 mb-3">
        <span class="badge bg-info text-dark mb-1">{{ e.kategori }}</span>
        <h5 class="fw-bold mb-1">{{ e.judul }}</h5>
        <p class="small text-muted mb-0">{{ e.isi[:180] }}...</p>
    </div>
    {% else %}
    <p class="text-muted small">Tidak ada esai yang cocok.</p>
    {% endfor %}
</div>
</body>
</html>
"""

HTML_BUKU_DETAIL = """
<!DOCTYPE html>
<html lang="id">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Detail Buku - Dede Suhendra</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
    <link href="https://fonts.googleapis.com/css2?family=Cinzel:wght@600;700&family=Plus+Jakarta+Sans:wght@400;600;700&display=swap" rel="stylesheet">
    <script src="https://cdn.jsdelivr.net/npm/marked/marked.min.js"></script>
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css">
    <style>""" + CSS_SHARED + """</style>
    """ + JS_THEME_SCRIPT + """
</head>
<body>
<div id="reading-progress"></div>

<button class="btn btn-mode-toggle" onclick="toggleModeInstan()" title="Ganti Mode Tampilan">
    <i class="fa-solid fa-moon" id="icon-mode"></i>
</button>

<a href="#" class="btn-to-top" id="btnScrollTop" title="Kembali ke Atas"><i class="fa-solid fa-arrow-up"></i></a>

<div class="container py-5" style="max-width:850px;">
    <div class="d-flex justify-content-between align-items-center mb-4 flex-wrap gap-2">
        <a href="/tema/{{ buku.tema_id }}" class="btn btn-custom-outline rounded-pill btn-sm px-4 shadow-sm fw-bold"><i class="fa-solid fa-arrow-left me-2"></i>Kembali ke Tema</a>
        <div class="d-flex gap-2 align-items-center flex-wrap">
            <a href="/export-buku/{{ buku.id }}" class="btn btn-outline-warning btn-sm rounded-pill fw-bold" title="TXT"><i class="fa-solid fa-file-arrow-down me-1"></i> TXT</a>
            <a href="/export-buku-docx/{{ buku.id }}" class="btn btn-outline-info btn-sm rounded-pill fw-bold" title="Word"><i class="fa-solid fa-file-word me-1"></i> DOCX</a>
            <a href="/export-buku-epub/{{ buku.id }}" class="btn btn-outline-success btn-sm rounded-pill fw-bold" title="E-Book"><i class="fa-solid fa-book me-1"></i> EPUB</a>
            <a href="/export-buku-pdf/{{ buku.id }}" class="btn btn-warning text-dark btn-sm rounded-pill fw-bold"><i class="fa-solid fa-file-pdf me-1"></i> PDF Buku</a>
            {% if is_admin %}
            <button type="button" class="btn btn-info btn-sm rounded-pill fw-bold text-dark" data-bs-toggle="modal" data-bs-target="#editBukuModal"><i class="fa-solid fa-pen-to-square me-1"></i> Edit Buku</button>
            <button type="button" class="btn btn-warning btn-sm rounded-pill fw-bold text-dark" data-bs-toggle="modal" data-bs-target="#modalUrutanBab"><i class="fa-solid fa-arrow-down-up-2-1 me-1"></i> Ubah Urutan Bab</button>
            {% endif %}
        </div>
    </div>
    
    <div class="card-gold p-4 mb-4 rounded-4">
        <span class="badge bg-warning text-dark mb-2 fw-bold"><i class="fa-solid fa-book me-1"></i> Naskah Buku</span>
        <h2 class="h3 fw-bold mb-1" style="font-family:'Cinzel',serif;">{{ buku.judul.upper() }}</h2>
        {% if buku.subjudul %}<p class="text-muted mb-2">{{ buku.subjudul }}</p>{% endif %}
        {% if buku.kutipan %}<blockquote class="blockquote small text-muted fst-italic mb-0">"{{ buku.kutipan }}"</blockquote>{% endif %}
    </div>

    {% if catatan_grouped %}
    <div class="toc-box">
        <h6 class="fw-bold mb-3 text-warning"><i class="fa-solid fa-list me-2"></i> Daftar Isi Buku</h6>
        {% for bagian, cats in catatan_grouped.items() %}
            <div class="toc-section-title"><i class="fa-solid fa-bookmark me-1 text-warning"></i> {{ bagian }}</div>
            <ul class="toc-list">
                {% for c in cats %}
                <li>
                    <a href="#bab-{{ c.id }}">
                        <i class="fa-solid fa-angle-right text-muted small"></i> {{ c.judul_bab }}
                    </a>
                </li>
                {% endfor %}
            </ul>
        {% endfor %}
    </div>
    {% endif %}

    {% if is_admin %}
    <div class="modal fade" id="modalUrutanBab" tabindex="-1" aria-hidden="true">
      <div class="modal-dialog modal-lg">
        <div class="modal-content bg-dark text-white border-secondary">
          <div class="modal-header border-secondary">
            <h5 class="modal-title fw-bold text-warning"><i class="fa-solid fa-arrow-down-up-2-1 me-2"></i>Atur Urutan Bab</h5>
            <button type="button" class="btn-close btn-close-white" data-bs-dismiss="modal"></button>
          </div>
          <form action="/ubah-urutan-bab/{{ buku.id }}" method="POST">
              <input type="hidden" name="csrf_token" value="{{ csrf_token() }}">
              <div class="modal-body" style="max-height:60vh; overflow-y:auto;">
                <p class="small text-muted mb-3">Ubah angka urutan bab di bawah ini (semakin kecil angkanya, semakin tinggi posisinya):</p>
                {% for cat in catatan_list %}
                <div class="d-flex align-items-center gap-2 mb-2 p-2 bg-secondary rounded">
                    <input type="hidden" name="catatan_ids[]" value="{{ cat.id }}">
                    <input type="number" name="urutan_vals[]" class="form-control form-control-sm text-center fw-bold" style="width:70px;" value="{{ cat.urutan }}" min="1">
                    <span class="badge bg-warning text-dark">{{ cat.bagian or 'Utama' }}</span>
                    <strong class="text-truncate">{{ cat.judul_bab }}</strong>
                </div>
                {% endfor %}
              </div>
              <div class="modal-footer border-secondary">
                <button type="button" class="btn btn-sm btn-outline-light" data-bs-dismiss="modal">Batal</button>
                <button type="submit" class="btn btn-sm btn-warning fw-bold text-dark">Simpan Urutan</button>
              </div>
          </form>
        </div>
      </div>
    </div>

    <div class="modal fade" id="editBukuModal" tabindex="-1" aria-hidden="true">
      <div class="modal-dialog">
        <div class="modal-content bg-dark text-white border-secondary">
          <div class="modal-header border-secondary">
            <h5 class="modal-title fw-bold">✏️ Edit Informasi Buku</h5>
            <button type="button" class="btn-close btn-close-white" data-bs-dismiss="modal"></button>
          </div>
          <form action="/edit-buku/{{ buku.id }}" method="POST">
              <input type="hidden" name="csrf_token" value="{{ csrf_token() }}">
              <div class="modal-body">
                <div class="mb-2">
                    <label class="small text-muted">Judul Buku</label>
                    <input type="text" name="judul" class="form-control bg-secondary text-white" value="{{ buku.judul }}" required>
                </div>
                <div class="mb-2">
                    <label class="small text-muted">Subjudul</label>
                    <input type="text" name="subjudul" class="form-control bg-secondary text-white" value="{{ buku.subjudul or '' }}">
                </div>
                <div class="mb-2">
                    <label class="small text-muted">Kutipan / Quote</label>
                    <textarea name="kutipan" class="form-control bg-secondary text-white" rows="2">{{ buku.kutipan or '' }}</textarea>
                </div>
              </div>
              <div class="modal-footer border-secondary">
                <button type="button" class="btn btn-sm btn-outline-light" data-bs-dismiss="modal">Batal</button>
                <button type="submit" class="btn btn-sm btn-info fw-bold text-dark">Simpan Perubahan</button>
              </div>
          </form>
        </div>
      </div>
    </div>
    
    <div class="card-gold p-3 mb-4 rounded-3">
        <h6 class="fw-bold text-success mb-2"><i class="fa-solid fa-file-circle-plus me-1"></i> Tambah Bab / Catatan Baru</h6>
        <form action="/tambah-catatan" method="POST">
            <input type="hidden" name="csrf_token" value="{{ csrf_token() }}">
            <input type="hidden" name="buku_id" value="{{ buku.id }}">
            <div class="row g-2 mb-2">
                <div class="col-md-3"><input type="text" name="bagian" class="form-control form-control-sm" placeholder="Kategori/Bagian"></div>
                <div class="col-md-7"><input type="text" name="judul_bab" class="form-control form-control-sm" placeholder="Judul Bab..." required></div>
                <div class="col-md-2"><input type="number" name="urutan" class="form-control form-control-sm" placeholder="Urutan (cth: 1)" value="1" min="1"></div>
            </div>
            <div class="mb-2"><textarea name="isi" class="form-control form-control-sm" rows="5" placeholder="Tulis isi catatan atau naskah bab (Markdown didukung)..." required></textarea></div>
            <button type="submit" class="btn btn-success btn-sm w-100 fw-bold">Simpan Bab</button>
        </form>
    </div>
    {% endif %}

    <h5 class="fw-bold mb-3" style="font-family:'Cinzel',serif;">Daftar Bab & Catatan:</h5>
    <div class="row g-3">
        {% for bagian, cats in catatan_grouped.items() %}
        <div class="col-12">
            <div class="section-header-card"><i class="fa-solid fa-bookmark me-2"></i> {{ bagian }}</div>
        </div>
        {% for cat in cats %}
        <div class="col-12" id="bab-{{ cat.id }}">
            <div class="card-gold p-4 position-relative">
                <div class="position-absolute top-0 end-0 p-3 d-flex align-items-center gap-1 flex-wrap justify-content-end" style="max-width: 320px;">
                    <span class="badge bg-dark text-warning border border-warning px-2 py-1 me-1">Urutan #{{ cat.urutan }}</span>
                    <button type="button" id="btnAudio-{{ cat.id }}" class="btn btn-sm btn-info text-dark rounded-pill fw-bold px-2 py-1 shadow-sm" onclick="playBabTTS({{ cat.id }}, '{{ cat.judul_bab | e }}')">
                        <i class="fa-solid fa-headphones me-1"></i> <span id="labelAudio-{{ cat.id }}">Dengar</span>
                    </button>

                    <form action="/tandai-baca/{{ cat.id }}" method="POST" class="mb-0">
                        <input type="hidden" name="csrf_token" value="{{ csrf_token() }}">
                        <button type="submit" class="btn btn-sm {% if penanda_aktif and penanda_aktif.catatan_id == cat.id %}btn-warning text-dark{% else %}btn-outline-secondary{% endif %} rounded-pill fw-bold px-2 py-1" title="Tandai posisi bacaan terakhir">
                            <i class="fa-solid fa-bookmark me-1"></i> {% if penanda_aktif and penanda_aktif.catatan_id == cat.id %}Dibaca{% else %}Tandai{% endif %}
                        </button>
                    </form>

                    {% if is_admin %}
                    <button type="button" class="btn btn-link text-info p-1" data-bs-toggle="modal" data-bs-target="#editBabModal{{ cat.id }}" title="Edit Bab">
                        <i class="fa-solid fa-pen-to-square fs-5"></i>
                    </button>
                    <form action="/toggle-favorit/{{ cat.id }}" method="POST" class="mb-0">
                        <input type="hidden" name="csrf_token" value="{{ csrf_token() }}">
                        <button type="submit" class="btn btn-link p-1 {% if cat.favorit %}text-warning{% else %}text-muted{% endif %}" title="Tandai Favorit">
                            <i class="fa-solid fa-star fs-5"></i>
                        </button>
                    </form>
                    <form action="/hapus-catatan/{{ cat.id }}/{{ buku.id }}" method="POST" class="mb-0" onsubmit="return confirm('Hapus bab ini?');">
                        <input type="hidden" name="csrf_token" value="{{ csrf_token() }}">
                        <button type="submit" class="btn btn-link text-danger p-1" title="Hapus Bab"><i class="fa-solid fa-trash fs-5"></i></button>
                    </form>
                    {% endif %}
                </div>

                {% if cat.bagian %}<span class="note-card-badge">{{ cat.bagian }}</span>{% endif %}
                <h4 class="note-card-title">{{ cat.judul_bab }}</h4>
                <div class="markdown-body text-main tts-isi-bab-{{ cat.id }}" id="content-catatan-{{ cat.id }}"></div>
                <textarea id="raw-catatan-{{ cat.id }}" style="display:none;">{{ cat.isi }}</textarea>
                <div class="text-muted small mt-3">Dibuat: {{ cat.dibuat_pada.strftime('%d %b %Y %H:%M') if cat.dibuat_pada else '-' }}</div>
            </div>
        </div>

        {% if is_admin %}
        <div class="modal fade" id="editBabModal{{ cat.id }}" tabindex="-1" aria-hidden="true">
          <div class="modal-dialog modal-lg">
            <div class="modal-content bg-dark text-white border-secondary">
              <div class="modal-header border-secondary">
                <h5 class="modal-title fw-bold">✏️ Edit Bab: {{ cat.judul_bab }}</h5>
                <button type="button" class="btn-close btn-close-white" data-bs-dismiss="modal"></button>
              </div>
              <form action="/edit-catatan/{{ cat.id }}" method="POST">
                  <input type="hidden" name="csrf_token" value="{{ csrf_token() }}">
                  <div class="modal-body">
                    <div class="row g-2 mb-2">
                        <div class="col-md-3">
                            <label class="small text-muted">Bagian</label>
                            <input type="text" name="bagian" class="form-control bg-secondary text-white" value="{{ cat.bagian or '' }}">
                        </div>
                        <div class="col-md-7">
                            <label class="small text-muted">Judul Bab</label>
                            <input type="text" name="judul_bab" class="form-control bg-secondary text-white" value="{{ cat.judul_bab }}" required>
                        </div>
                        <div class="col-md-2">
                            <label class="small text-muted">Urutan</label>
                            <input type="number" name="urutan" class="form-control bg-secondary text-white" value="{{ cat.urutan }}" min="1">
                        </div>
                    </div>
                    <div class="mb-2">
                        <label class="small text-muted">Isi Naskah (Markdown)</label>
                        <textarea name="isi" class="form-control bg-secondary text-white" rows="10" required>{{ cat.isi }}</textarea>
                    </div>
                  </div>
                  <div class="modal-footer border-secondary">
                    <button type="button" class="btn btn-sm btn-outline-light" data-bs-dismiss="modal">Batal</button>
                    <button type="submit" class="btn btn-sm btn-success fw-bold">Simpan Perubahan</button>
                  </div>
              </form>
            </div>
          </div>
        </div>
        {% endif %}
        {% endfor %}
        
        {% else %}
        <div class="text-center py-4 card-gold rounded-4">
            <p class="text-muted mb-0">Belum ada catatan bab dalam buku ini.</p>
        </div>
        {% endfor %}
    </div>
</div>

<script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/js/bootstrap.bundle.min.js"></script>
<script>
document.addEventListener("DOMContentLoaded", function(){
    marked.use({ gfm: true, breaks: true });
    {% for cat in catatan_list %}
        var rawText = document.getElementById('raw-catatan-{{ cat.id }}').value;
        document.getElementById('content-catatan-{{ cat.id }}').innerHTML = marked.parse(rawText);
    {% endfor %}
});

let synth = window.speechSynthesis;
let activeUtteranceId = null;

function playBabTTS(id, judulBab) {
    if (!synth) {
        alert("Browser Anda tidak mendukung fitur suara (Text-to-Speech).");
        return;
    }

    const btn = document.getElementById('btnAudio-' + id);
    const label = document.getElementById('labelAudio-' + id);

    if (synth.speaking && activeUtteranceId === id) {
        synth.cancel();
        resetTombolTTS(id);
        activeUtteranceId = null;
        return;
    }

    if (synth.speaking) {
        synth.cancel();
        if (activeUtteranceId !== null) {
            resetTombolTTS(activeUtteranceId);
        }
    }

    const contentEl = document.querySelector('.tts-isi-bab-' + id);
    if (!contentEl) return;

    let textToRead = "Judul bab: " + judulBab + ". " + contentEl.innerText;
    textToRead = textToRead.replace(/[«»“”"''*#_`~]/g, '');

    let utterance = new SpeechSynthesisUtterance(textToRead);
    utterance.lang = 'id-ID';
    utterance.rate = 1.0;

    utterance.onstart = function() {
        activeUtteranceId = id;
        label.innerText = "Stop ⏹";
        btn.className = "btn btn-sm btn-danger text-white rounded-pill fw-bold px-2 py-1 shadow-sm";
    };

    utterance.onend = function() {
        resetTombolTTS(id);
        activeUtteranceId = null;
    };

    utterance.onerror = function() {
        resetTombolTTS(id);
        activeUtteranceId = null;
    };

    synth.speak(utterance);
}

function resetTombolTTS(id) {
    const btn = document.getElementById('btnAudio-' + id);
    const label = document.getElementById('labelAudio-' + id);
    if (btn && label) {
        label.innerText = "Dengar";
        btn.className = "btn btn-sm btn-info text-dark rounded-pill fw-bold px-2 py-1 shadow-sm";
    }
}
</script>
</body>
</html>
"""

HTML_FAVORIT = """<!DOCTYPE html><html lang="id"><head><meta charset="UTF-8"><title>Favorit</title><link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet"><link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css"><style>""" + CSS_SHARED + """</style>""" + JS_THEME_SCRIPT + """</head><body><button class="btn btn-mode-toggle" onclick="toggleModeInstan()"><i class="fa-solid fa-moon" id="icon-mode"></i></button><div class="container py-4" style="max-width:750px;"><a href="/" class="btn btn-custom-outline btn-sm mb-4">&larr; Kembali</a><h4 class="fw-bold mb-3 text-warning"><i class="fa-solid fa-star me-2"></i> Koleksi Favorit</h4><div class="d-flex flex-column gap-3">{% for c in catatan_favorit %}<div class="card-gold p-4"><div class="d-flex justify-content-between align-items-center mb-2"><span class="badge bg-warning text-dark">Buku ID: {{ c.buku_id }}</span><a href="/buku/{{ c.buku_id }}" class="btn btn-outline-warning btn-sm rounded-pill fw-bold">Buka Buku &rarr;</a></div><h5 class="fw-bold mb-2">{{ c.judul_bab }}</h5><p class="small text-muted mb-0">{{ c.isi[:140] }}...</p></div>{% else %}<div class="text-center py-5 card-gold rounded-4"><p class="text-muted mb-0">Belum ada bab favorit.</p></div>{% endfor %}</div></div></body></html>"""
HTML_PDF_VIEWER = """<!DOCTYPE html><html lang="id"><head><meta charset="UTF-8"><title>PDF Reader</title><link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet"><link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css"><style>""" + CSS_SHARED + """.pdf-frame-wrapper{position:relative;width:100%;height:75vh;border-radius:14px;overflow:hidden;border:1px solid var(--border-color);}iframe{width:100%;height:100%;border:none;}</style>""" + JS_THEME_SCRIPT + """</head><body><button class="btn btn-mode-toggle" onclick="toggleModeInstan()"><i class="fa-solid fa-moon" id="icon-mode"></i></button><div class="container py-4" style="max-width:920px;"><a href="/" class="btn btn-custom-outline btn-sm mb-3">&larr; Kembali</a>{% if is_admin %}<div class="card-gold p-3 mb-4 rounded-3"><form action="/upload-pdf" method="POST" enctype="multipart/form-data" class="row g-2"><input type="hidden" name="csrf_token" value="{{ csrf_token() }}"><div class="col-md-9"><input type="file" name="file_pdf" class="form-control form-control-sm" accept=".pdf" required></div><div class="col-md-3"><button type="submit" class="btn btn-success btn-sm w-100 fw-bold">Upload PDF</button></div></form></div>{% endif %}<div class="card-gold p-3 mb-3 d-flex justify-content-between align-items-center flex-wrap gap-2"><h4 class="h5 fw-bold mb-0 text-warning">{{ nama_file }}</h4>{% if nama_file != "Pilih file PDF di bawah" %}<a href="/file-pdf/{{ nama_file }}" download class="btn btn-outline-warning btn-sm rounded-pill fw-bold"><i class="fa-solid fa-download me-1"></i> Download</a>{% endif %}</div><div class="card-gold p-3 mb-3"><div class="d-flex flex-wrap gap-2">{% for f in daftar_file %}<a href="/baca-pdf?nama={{ f }}" class="btn btn-sm {% if f == nama_file %}btn-warning text-dark{% else %}btn-custom-outline{% endif %} rounded-pill fw-bold"><i class="fa-solid fa-file-pdf me-1"></i> {{ f }}</a>{% endfor %}</div></div>{% if url_pdf %}<div class="pdf-frame-wrapper card-gold"><iframe src="https://mozilla.github.io/pdf.js/web/viewer.html?file={{ url_pdf }}"></iframe></div>{% endif %}</div></body></html>"""
HTML_ESAI_PENULIS = """<!DOCTYPE html><html lang="id"><head><meta charset="UTF-8"><title>Esai</title><link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet"><script src="https://cdn.jsdelivr.net/npm/marked/marked.min.js"></script><link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css"><style>""" + CSS_SHARED + """</style>""" + JS_THEME_SCRIPT + """</head><body><button class="btn btn-mode-toggle" onclick="toggleModeInstan()"><i class="fa-solid fa-moon" id="icon-mode"></i></button><div class="container py-4" style="max-width:760px;"><a href="/" class="btn btn-custom-outline btn-sm mb-4">&larr; Kembali</a><div class="card-gold p-4 rounded-4 mb-4"><h2 class="h3 fw-bold text-warning mb-1">✍️ Jurnal & Esai Bebas</h2></div>{% if is_admin %}<div class="card-gold p-3 mb-4 rounded-3"><form action="/tambah-esai" method="POST"><input type="hidden" name="csrf_token" value="{{ csrf_token() }}"><input type="text" name="judul" class="form-control form-control-sm mb-2" placeholder="Judul..." required><input type="text" name="kategori" class="form-control form-control-sm mb-2" placeholder="Kategori..."><textarea name="isi" class="form-control form-control-sm mb-2" rows="4" placeholder="Isi..." required></textarea><button type="submit" class="btn btn-info btn-sm w-100 fw-bold">Terbitkan</button></form></div>{% endif %}{% for e in esai_list %}<div class="card-gold p-4 mb-3"><div class="d-flex justify-content-between mb-2"><span class="badge bg-info text-dark">{{ e.kategori }}</span><a href="/cetak-esai-pdf/{{ e.id }}" class="btn btn-outline-warning btn-sm rounded-pill fw-bold"><i class="fa-solid fa-file-pdf"></i> PDF</a></div><h4 class="text-warning fw-bold mb-3">{{ e.judul }}</h4><div class="markdown-body" id="content-esai-{{ e.id }}"></div><textarea id="raw-esai-{{ e.id }}" style="display:none;">{{ e.isi }}</textarea></div>{% endfor %}</div><script>document.addEventListener("DOMContentLoaded", function(){marked.use({ gfm: true, breaks: true });{% for e in esai_list %}document.getElementById('content-esai-{{ e.id }}').innerHTML = marked.parse(document.getElementById('raw-esai-{{ e.id }}').value);{% endfor %});</script></body></html>"""

HTML_TEMA = """
<!DOCTYPE html>
<html lang="id">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Tema: {{ tema.nama }}</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css">
    <style>""" + CSS_SHARED + """</style>
    """ + JS_THEME_SCRIPT + """
</head>
<body>
<button class="btn btn-mode-toggle" onclick="toggleModeInstan()"><i class="fa-solid fa-moon" id="icon-mode"></i></button>
<div class="container py-5" style="max-width:850px;">
    <a href="/" class="btn btn-custom-outline rounded-pill btn-sm px-4 mb-4 fw-bold">&larr; Kembali</a>
    <div class="card-gold p-4 mb-4 rounded-4">
        <h2 class="h3 fw-bold mb-1">Tema: {{ tema.nama }}</h2>
    </div>
    {% if is_admin %}
    <div class="card-gold p-3 mb-4 rounded-3">
        <form action="/tambah-buku" method="POST" class="row g-2">
            <input type="hidden" name="csrf_token" value="{{ csrf_token() }}">
            <input type="hidden" name="tema_id" value="{{ tema.id }}">
            <div class="col-md-6"><input type="text" name="judul" class="form-control form-control-sm" placeholder="Judul Buku..." required></div>
            <div class="col-md-6"><input type="text" name="subjudul" class="form-control form-control-sm" placeholder="Subjudul"></div>
            <div class="col-12"><textarea name="kutipan" class="form-control form-control-sm" rows="2" placeholder="Kutipan..."></textarea></div>
            <div class="col-12"><button type="submit" class="btn btn-success btn-sm w-100 fw-bold">Simpan Buku</button></div>
        </form>
    </div>
    {% endif %}
    <div class="row g-3">
        {% for buku in buku_list %}
        <div class="col-12">
            <div class="card-gold p-4 d-flex justify-content-between align-items-center flex-wrap gap-2">
                <a href="/buku/{{ buku.id }}" class="text-decoration-none flex-grow-1">
                    <h4 class="h5 mb-1 fw-bold text-warning">{{ buku.judul.upper() }}</h4>
                    {% if buku.subjudul %}<p class="text-muted small mb-0">{{ buku.subjudul }}</p>{% endif %}
                </a>
                <div class="d-flex align-items-center gap-2">
                    <a href="/export-buku/{{ buku.id }}" class="btn btn-outline-warning btn-sm rounded-pill fw-bold" title="TXT"><i class="fa-solid fa-file-arrow-down"></i></a>
                    <a href="/export-buku-pdf/{{ buku.id }}" class="btn btn-warning text-dark btn-sm rounded-pill fw-bold" title="PDF"><i class="fa-solid fa-file-pdf"></i></a>
                    {% if is_admin %}
                    <button type="button" class="btn btn-outline-info btn-sm rounded-pill fw-bold" data-bs-toggle="modal" data-bs-target="#pindahTemaModal{{ buku.id }}" title="Pindah Tema"><i class="fa-solid fa-folder-tree"></i></button>
                    <form action="/hapus-buku/{{ buku.id }}/{{ tema.id }}" method="POST" class="mb-0" onsubmit="return confirm('Hapus buku ini beserta seluruh isinya?');">
                        <input type="hidden" name="csrf_token" value="{{ csrf_token() }}">
                        <button type="submit" class="btn btn-outline-danger btn-sm rounded-pill fw-bold" title="Hapus Buku"><i class="fa-solid fa-trash"></i></button>
                    </form>
                    {% endif %}
                </div>
            </div>
        </div>

        {% if is_admin %}
        <div class="modal fade" id="pindahTemaModal{{ buku.id }}" tabindex="-1" aria-hidden="true">
          <div class="modal-dialog">
            <div class="modal-content bg-dark text-white border-secondary">
              <div class="modal-header border-secondary">
                <h5 class="modal-title fw-bold">📂 Pindah Buku ke Tema Lain</h5>
                <button type="button" class="btn-close btn-close-white" data-bs-dismiss="modal"></button>
              </div>
              <form action="/pindah-tema-buku/{{ buku.id }}" method="POST">
                  <input type="hidden" name="csrf_token" value="{{ csrf_token() }}">
                  <div class="modal-body">
                    <p class="small text-muted">Pilih tema tujuan untuk buku: <strong>{{ buku.judul }}</strong></p>
                    <select name="tema_baru_id" class="form-select bg-secondary text-white border-secondary" required>
                        {% for t_all in semua_tema %}
                        <option value="{{ t_all.id }}" {% if t_all.id == tema.id %}selected{% endif %}>{{ t_all.nama }}</option>
                        {% endfor %}
                    </select>
                  </div>
                  <div class="modal-footer border-secondary">
                    <button type="button" class="btn btn-sm btn-outline-light" data-bs-dismiss="modal">Batal</button>
                    <button type="submit" class="btn btn-sm btn-info fw-bold text-dark">Pindahkan Buku</button>
                  </div>
              </form>
            </div>
          </div>
        </div>
        {% endif %}
        {% endfor %}
    </div>
</div>
<script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/js/bootstrap.bundle.min.js"></script>
</body>
</html>
"""

HTML_PENULIS = """
<!DOCTYPE html>
<html lang="id">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Profil Penulis - Dede Suhendra</title>
    <link href="https://fonts.googleapis.com/css2?family=Cinzel:wght@600;700&family=Plus+Jakarta+Sans:wght@400;600;700;800&display=swap" rel="stylesheet">
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css">
    <style>""" + CSS_SHARED + """</style>
    """ + JS_THEME_SCRIPT + """
</head>
<body>
<button class="btn btn-mode-toggle" onclick="toggleModeInstan()" title="Ganti Mode Tampilan">
    <i class="fa-solid fa-moon" id="icon-mode"></i>
</button>

<div class="container py-5" style="max-width:760px;">
    <a href="/" class="btn btn-custom-outline btn-sm mb-4 rounded-pill px-3 fw-bold">&larr; Kembali ke Utama</a>
    
    <div class="card-gold p-4 mb-4 rounded-4" style="display:flex; gap:20px; align-items:center; flex-wrap:wrap;">
        <img src="/profile.jpg" onerror="this.src='https://cdn-icons-png.flaticon.com/512/3135/3135715.png'" alt="Dede Suhendra" style="width:140px; height:140px; border-radius:50%; object-fit:cover; border:3px solid #b38728; margin:0 auto; display:block;">
        <div style="flex:1; min-width:250px;">
            <h2 class="h4 fw-bold text-warning mb-2" style="font-family:'Cinzel',serif;">👨‍💻 Profil Penulis</h2>
            <p class="small mb-2">Selamat datang di ruang pustaka pribadi karya dan catatan saya. Nama saya <strong>Dede Suhendra</strong>, lahir 8 Juli 2001, dari Subang.</p>
            <p class="small text-muted mb-0">Dokumentasi pemikiran, perjalanan belajar, riset harian, serta modul pembelajaran yang disusun terstruktur.</p>
        </div>
    </div>
    
    <div class="card-gold p-4 mb-4 rounded-4">
        <h4 class="h5 fw-bold text-warning mb-3">✍️ Perjuangan & Latar Belakang Penulisan</h4>
        <p class="small mb-2">Setiap tulisan lahir dari proses yang tidak instan. Di tengah padatnya aktivitas harian, setiap sisa waktu luang dimanfaatkan untuk tetap konsisten menulis dan mendokumentasikan ilmu.</p>
        <p class="small text-muted mb-0">Bagi saya, menulis bukan sekadar merangkai kata, melainkan bentuk pengikatan ilmu dan sarana merefleksikan pembelajaran hidup agar bermanfaat secara luas dan berkelanjutan.</p>
    </div>
    
    <div class="card-gold p-4 mb-4 rounded-4">
        <h4 class="h5 fw-bold text-warning mb-3">📜 Riwayat Pendidikan & Pengalaman</h4>
        <h6 class="fw-bold text-warning small mb-2">🎓 Pendidikan:</h6>
        <div style="border-left:2px solid var(--border-color); padding-left:15px; margin-bottom:15px;" class="small">
            <div class="mb-2"><strong>2013:</strong> SDN Sindang Laut II (Lulus SD)</div>
            <div class="mb-2"><strong>2013–2015:</strong> Ponpes Madinatul Musthofa</div>
            <div class="mb-2"><strong>2015–2016:</strong> Pondok Tahfidz Qur'an (Fokus Hafalan)</div>
            <div class="mb-2"><strong>2016–2019:</strong> Ponpes Madinatul Musthofa</div>
            <div class="mb-2"><strong>2019–2022:</strong> Pondok Modern Darussalam Gontor (KMI)</div>
            <div class="mb-2"><strong>2022–2023:</strong> Pengabdian Gontor & UNIDA Gontor</div>
            <div><strong>2023–2025:</strong> Pengajar Ponpes & STISQ AL-IHYA Subang</div>
        </div>
        <h6 class="fw-bold text-warning small mb-2">💼 Pengalaman Kerja & Khidmat:</h6>
        <div style="border-left:2px solid var(--border-color); padding-left:15px;" class="small">
            <div class="mb-2"><strong>2025:</strong> Gudang Shopee Tangerang (Logistik)</div>
            <div class="mb-2"><strong>2025:</strong> Karyawan Fotokopi Jakarta Pusat</div>
            <div class="mb-2"><strong>2025:</strong> Barista & Chef Bogor</div>
            <div><strong>Sekarang:</strong> Imam, Muadzin & Pengajar Al-Qur'an Tangerang</div>
        </div>
    </div>
    
    <div class="card-gold p-4 mb-4 rounded-4">
        <h4 class="h5 fw-bold text-warning mb-3">🎯 Visi & Misi Penulisan</h4>
        <p class="small mb-2"><strong>Visi:</strong> Menjadikan dokumentasi catatan pribadi sebagai sarana pengikat ilmu, pengembangan diri berkelanjutan, dan ladang manfaat terstruktur.</p>
        <p class="small fw-bold mb-1">Misi:</p>
        <ul class="small text-muted ps-3 mb-0">
            <li>Memanfaatkan setiap sisa waktu luang secara produktif untuk merangkai karya tulis dan modul bermanfaat.</li>
            <li>Mendokumentasikan pemahaman keagamaan, riset harian, dan keterampilan operasional secara rapi dan terbuka.</li>
            <li>Terus belajar dan memberikan dampak positif bagi santri, jamaah masjid, serta lingkungan sekitar.</li>
        </ul>
    </div>
    
    <div class="card-gold p-4 rounded-4">
        <h4 class="h5 fw-bold text-warning mb-3">🙏 Apresiasi & Rasa Syukur</h4>
        <p class="small text-muted mb-3">Rasa syukur dan terima kasih kepada orang-orang terkasih yang menjadi sumber kekuatan, doa, dan inspirasi:</p>
        <div class="row g-3">
            <div class="col-md-6"><div class="p-3 border rounded-3 h-100" style="border-color:var(--border-color)!important;"><strong class="text-warning">👨‍👦 Bapak Khairudin</strong><p class="small text-muted mb-0 mt-1">Doa, kerja keras, dan bimbingan tanpa henti.</p></div></div>
            <div class="col-md-6"><div class="p-3 border rounded-3 h-100" style="border-color:var(--border-color)!important;"><strong class="text-warning">💐 Ibu Sumini (Almarhumah)</strong><p class="small text-muted mb-0 mt-1">Semoga Allah mengampuni dan menempatkan di tempat terbaik.</p></div></div>
            <div class="col-md-6"><div class="p-3 border rounded-3 h-100" style="border-color:var(--border-color)!important;"><strong class="text-warning">👫 Siti Aisyah & Muhammad Naimul Ilmi</strong><p class="small text-muted mb-0 mt-1">Adik-adik tersayang, kebanggaan dan penyemangat.</p></div></div>
            <div class="col-md-6"><div class="p-3 border rounded-3 h-100" style="border-color:var(--border-color)!important;"><strong class="text-warning">👦 Muhammad Aji</strong><p class="small text-muted mb-0 mt-1">Kakak tercinta atas kebersamaan dan dukungan.</p></div></div>
            <div class="col-md-6"><div class="p-3 border rounded-3 h-100" style="border-color:var(--border-color)!important;"><strong class="text-warning">❤️ Sri Nur Safitri</strong><p class="small text-muted mb-0 mt-1">Perhatian, dorongan semangat, dan pendamping setia.</p></div></div>
            <div class="col-md-6"><div class="p-3 border rounded-3 h-100" style="border-color:var(--border-color)!important;"><strong class="text-warning">🤝 Sahabat & Kolega</strong><p class="small text-muted mb-0 mt-1">Semua yang telah mendukung dan mendoakan.</p></div></div>
        </div>
    </div>
</div>
</body>
</html>
"""

HTML_STATISTIK = """<!DOCTYPE html><html lang="id"><head><meta charset="UTF-8"><title>Statistik</title><link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet"><link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css"><style>""" + CSS_SHARED + """.stat-number{font-size:42px;font-weight:800;color:#b38728;}</style>""" + JS_THEME_SCRIPT + """</head><body><button class="btn btn-mode-toggle" onclick="toggleModeInstan()"><i class="fa-solid fa-moon" id="icon-mode"></i></button><div class="container py-4" style="max-width:600px;"><a href="/" class="btn btn-custom-outline btn-sm mb-4">&larr; Kembali</a><h3 class="mb-4 text-center fw-bold">📊 Statistik</h3><div class="row g-3"><div class="col-6"><div class="card-gold text-center p-4"><div class="stat-number">{{ total_buku }}</div><div class="text-muted">Buku</div></div></div><div class="col-6"><div class="card-gold text-center p-4"><div class="stat-number">{{ total_catatan }}</div><div class="text-muted">Bab</div></div></div></div></div></body></html>"""
HTML_TONG_SAMPAH = """<!DOCTYPE html><html lang="id"><head><meta charset="UTF-8"><title>Tong Sampah</title><link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet"><link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css"><style>""" + CSS_SHARED + """</style>""" + JS_THEME_SCRIPT + """</head><body><button class="btn btn-mode-toggle" onclick="toggleModeInstan()"><i class="fa-solid fa-moon" id="icon-mode"></i></button><div class="container py-4" style="max-width:700px;"><a href="/" class="btn btn-custom-outline btn-sm mb-4">&larr; Kembali</a><h4 class="mb-4 fw-bold">🗑️ Tong Sampah</h4>{% for item in sampah_list %}<div class="card-gold p-3 mb-2 d-flex justify-content-between align-items-center"><div><span class="badge bg-secondary me-2">{{ item.tipe }}</span><span>{{ item.data_json }}</span></div><div class="d-flex gap-2"><a href="/pulihkan/{{ item.id }}" class="btn btn-sm btn-success"><i class="fa-solid fa-rotate-left"></i></a><a href="/hapus-permanen/{{ item.id }}" class="btn btn-sm btn-danger"><i class="fa-solid fa-trash-xmark"></i></a></div></div>{% else %}<div class="text-center py-5 card-gold rounded-4"><p class="text-muted">Kosong.</p></div>{% endfor %}</div></body></html>"""
HTML_LOG = """<!DOCTYPE html><html lang="id"><head><meta charset="UTF-8"><title>Log</title><link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet"><link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css"><style>""" + CSS_SHARED + """</style>""" + JS_THEME_SCRIPT + """</head><body><button class="btn btn-mode-toggle" onclick="toggleModeInstan()"><i class="fa-solid fa-moon" id="icon-mode"></i></button><div class="container py-4" style="max-width:750px;"><a href="/" class="btn btn-custom-outline btn-sm mb-4">&larr; Kembali</a><h4 class="fw-bold mb-3 text-warning">Riwayat Log</h4><div class="d-flex flex-column gap-2">{% for l in logs %}<div class="card-gold p-3"><span class="badge bg-warning text-dark fw-bold mb-1">{{ l.aksi }}</span><p class="mb-0 small">{{ l.keterangan }}</p></div>{% endfor %}</div></div></body></html>"""
HTML_LOGIN = """<!DOCTYPE html><html lang="id"><head><meta charset="UTF-8"><title>Login</title><link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet"><style>body{background:#0f172a;color:#f8fafc;display:flex;justify-content:center;align-items:center;min-height:100vh;}.login-card{background:#1e293b;border:1px solid #334155;padding:30px;border-radius:12px;width:100%;max-width:400px;}</style></head><body><div class="login-card shadow-lg"><h3 class="text-info text-center fw-bold mb-3">🔑 Login Admin</h3>{% if error %}<div class="alert alert-danger py-2 small text-center">{{ error }}</div>{% endif %}<form action="/login" method="POST"><input type="hidden" name="csrf_token" value="{{ csrf_token() }}"><input type="text" name="username" class="form-control mb-3 bg-dark text-white border-secondary" placeholder="Username" required><input type="password" name="password" class="form-control mb-3 bg-dark text-white border-secondary" placeholder="Password" required><button type="submit" class="btn btn-info w-100 fw-bold">Masuk</button></form><div class="text-center mt-3"><a href="/" class="text-muted text-decoration-none small">&larr; Kembali</a></div></div></body></html>"""

@app.route('/')
def index():
    is_admin = session.get('is_admin')
    tema_list = Tema.query.order_by(Tema.id.asc()).all()
    jumlah_esai = EsaiPenulis.query.count()
    penanda_list = PenandaBaca.query.order_by(PenandaBaca.waktu_baca.desc()).limit(1).all()
    return render_template_string(HTML_INDEX, tema_list=tema_list, is_admin=is_admin, jumlah_esai=jumlah_esai, penanda_list=penanda_list)

@app.route('/cari')
def cari_naskah():
    query = request.args.get('q', '').strip()
    hasil_catatan = []
    hasil_esai = []

    if query:
        hasil_catatan = Catatan.query.filter(
            (Catatan.judul_bab.ilike(f'%{query}%')) | 
            (Catatan.isi.ilike(f'%{query}%'))
        ).all()

        hasil_esai = EsaiPenulis.query.filter(
            (EsaiPenulis.judul.ilike(f'%{query}%')) | 
            (EsaiPenulis.isi.ilike(f'%{query}%'))
        ).all()

    return render_template_string(HTML_HASIL_CARI, query=query, hasil_catatan=hasil_catatan, hasil_esai=hasil_esai)

@app.route('/baca-pdf')
def baca_pdf():
    is_admin = session.get('is_admin')
    daftar_file = [f for f in os.listdir(UPLOAD_PDF_FOLDER) if f.endswith('.pdf')]
    nama_pilihan = request.args.get('nama')
    if not nama_pilihan and daftar_file:
        nama_pilihan = daftar_file[0]
    url_pdf = f"/file-pdf/{nama_pilihan}" if nama_pilihan in daftar_file else None
    return render_template_string(HTML_PDF_VIEWER, url_pdf=url_pdf, nama_file=nama_pilihan or "Pilih file PDF di bawah", daftar_file=daftar_file, is_admin=is_admin)

@app.route('/upload-pdf', methods=['POST'])
def upload_pdf():
    if not session.get('is_admin'): return "Akses Ditolak", 403
    file = request.files.get('file_pdf')
    if file and file.filename.endswith('.pdf'):
        filename = re.sub(r'[^a-zA-Z0-9_.-]', '_', file.filename)
        file.save(os.path.join(UPLOAD_PDF_FOLDER, filename))
        catat_log("UPLOAD PDF", f"Upload: {filename}")
    return redirect('/baca-pdf')

@app.route('/file-pdf/<path:filename>')
def serve_uploaded_pdf(filename):
    return send_from_directory(UPLOAD_PDF_FOLDER, filename)

@app.route('/penulis')
def tentang_penulis(): return render_template_string(HTML_PENULIS)
@app.route('/catatan-penulis')
def catatan_penulis(): return render_template_string(HTML_ESAI_PENULIS, esai_list=EsaiPenulis.query.order_by(EsaiPenulis.id.desc()).all(), is_admin=session.get('is_admin'))
@app.route('/favorit')
def halaman_favorit(): return render_template_string(HTML_FAVORIT, catatan_favorit=Catatan.query.filter_by(favorit=True).all())

@app.route('/toggle-favorit/<int:catatan_id>', methods=['POST'])
def toggle_favorit(catatan_id):
    if not session.get('is_admin'): return "Akses Ditolak", 403
    catatan = Catatan.query.get_or_404(catatan_id)
    catatan.favorit = not catatan.favorit
    db.session.commit()
    return redirect(f'/buku/{catatan.buku_id}')

@app.route('/tandai-baca/<int:catatan_id>', methods=['POST'])
def tandai_baca(catatan_id):
    catatan = Catatan.query.get_or_404(catatan_id)
    PenandaBaca.query.filter_by(buku_id=catatan.buku_id).delete()
    db.session.add(PenandaBaca(buku_id=catatan.buku_id, catatan_id=catatan.id))
    db.session.commit()
    return redirect(f'/buku/{catatan.buku_id}')

@app.route('/cetak-esai-pdf/<int:esai_id>')
def cetak_esai_pdf(esai_id):
    esai = EsaiPenulis.query.get_or_404(esai_id)
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    story = [Paragraph(f"<b>{escape(esai.judul)}</b>", styles['Heading1']), Paragraph(bersihkan_teks_pdf(esai.isi), styles['Normal'])]
    doc.build(story)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name=f"esai_{esai.id}.pdf", mimetype='application/pdf')

@app.route('/tambah-esai', methods=['POST'])
def tambah_esai():
    if not session.get('is_admin'): return "Akses Ditolak", 403
    db.session.add(EsaiPenulis(judul=request.form.get('judul', '').strip(), kategori=request.form.get('kategori', 'Refleksi').strip(), isi=request.form.get('isi', '').strip()))
    db.session.commit()
    return redirect('/catatan-penulis')

@app.route('/tema/<int:tema_id>')
def detail_tema(tema_id):
    tema = Tema.query.get_or_404(tema_id)
    buku_list = Buku.query.filter_by(tema_id=tema_id).all()
    semua_tema = Tema.query.order_by(Tema.nama.asc()).all()
    return render_template_string(HTML_TEMA, tema=tema, buku_list=buku_list, semua_tema=semua_tema, is_admin=session.get('is_admin'))

@app.route('/tambah-tema', methods=['POST'])
def tambah_tema():
    if not session.get('is_admin'): return "Akses Ditolak", 403
    nama = request.form.get('nama', '').strip()
    if nama:
        db.session.add(Tema(nama=nama))
        db.session.commit()
    return redirect('/')

@app.route('/hapus-tema/<int:tema_id>', methods=['POST'])
def hapus_tema(tema_id):
    if not session.get('is_admin'): return "Akses Ditolak", 403
    tema = Tema.query.get(tema_id)
    if tema:
        db.session.delete(tema)
        db.session.commit()
    return redirect('/')

@app.route('/tambah-buku', methods=['POST'])
def tambah_buku():
    if not session.get('is_admin'): return "Akses Ditolak", 403
    judul = request.form.get('judul', '').strip()
    tema_id = request.form.get('tema_id')
    if judul and tema_id:
        db.session.add(Buku(judul=judul, subjudul=request.form.get('subjudul', '').strip(), tema_id=int(tema_id), kutipan=request.form.get('kutipan', '').strip()))
        db.session.commit()
    return redirect(f'/tema/{tema_id}')

@app.route('/edit-buku/<int:buku_id>', methods=['POST'])
def edit_buku(buku_id):
    if not session.get('is_admin'): return "Akses Ditolak", 403
    buku = Buku.query.get_or_404(buku_id)
    buku.judul = request.form.get('judul', '').strip()
    buku.subjudul = request.form.get('subjudul', '').strip()
    buku.kutipan = request.form.get('kutipan', '').strip()
    db.session.commit()
    catat_log("EDIT BUKU", f"Memperbarui buku: {buku.judul}")
    return redirect(f'/buku/{buku.id}')

@app.route('/pindah-tema-buku/<int:buku_id>', methods=['POST'])
def pindah_tema_buku(buku_id):
    if not session.get('is_admin'): return "Akses Ditolak", 403
    buku = Buku.query.get_or_404(buku_id)
    tema_baru_id = request.form.get('tema_baru_id')
    if tema_baru_id:
        buku.tema_id = int(tema_baru_id)
        db.session.commit()
        catat_log("PINDAH TEMA", f"Memindahkan buku '{buku.judul}' ke tema ID {tema_baru_id}")
    return redirect(f'/tema/{buku.tema_id}')

@app.route('/hapus-buku/<int:buku_id>/<int:tema_id>', methods=['POST'])
def hapus_buku(buku_id, tema_id):
    if not session.get('is_admin'): return "Akses Ditolak", 403
    buku = Buku.query.get(buku_id)
    if buku:
        masukkan_sampah('buku', {'judul': buku.judul, 'subjudul': buku.subjudul})
        db.session.delete(buku)
        db.session.commit()
        catat_log("HAPUS BUKU", f"Menghapus buku: {buku.judul}")
    return redirect(f'/tema/{tema_id}')

@app.route('/buku/<int:buku_id>')
def detail_buku(buku_id):
    buku = Buku.query.get_or_404(buku_id)
    catatan_list = Catatan.query.filter_by(buku_id=buku_id).order_by(Catatan.urutan.asc(), Catatan.id.asc()).all()
    penanda_aktif = PenandaBaca.query.filter_by(buku_id=buku_id).first()
    
    # Pengelompokan Bab sesuai Bagian / Kategori untuk Daftar Isi & Tampilan Kartu
    catatan_grouped = defaultdict(list)
    for c in catatan_list:
        bagian_key = c.bagian if c.bagian else 'Bagian Utama'
        catatan_grouped[bagian_key].append(c)

    return render_template_string(HTML_BUKU_DETAIL, buku=buku, catatan_list=catatan_list, catatan_grouped=dict(catatan_grouped), is_admin=session.get('is_admin'), penanda_aktif=penanda_aktif)

@app.route('/ubah-urutan-bab/<int:buku_id>', methods=['POST'])
def ubah_urutan_bab(buku_id):
    if not session.get('is_admin'): return "Akses Ditolak", 403
    catatan_ids = request.form.getlist('catatan_ids[]')
    urutan_vals = request.form.getlist('urutan_vals[]')
    
    for c_id, u_val in zip(catatan_ids, urutan_vals):
        catatan = Catatan.query.get(c_id)
        if catatan:
            try:
                catatan.urutan = int(u_val)
            except ValueError:
                pass
    db.session.commit()
    catat_log("URUTAN BAB", f"Mengubah urutan bab buku ID: {buku_id}")
    return redirect(f'/buku/{buku_id}')

@app.route('/tambah-catatan', methods=['POST'])
def tambah_catatan():
    if not session.get('is_admin'): return "Akses Ditolak", 403
    buku_id = request.form.get('buku_id')
    judul_bab = request.form.get('judul_bab', '').strip()
    isi = request.form.get('isi', '').strip()
    bagian = request.form.get('bagian', '').strip()
    urutan_val = request.form.get('urutan', 1)

    try:
        urutan_val = int(urutan_val)
    except ValueError:
        urutan_val = 1

    if buku_id and judul_bab and isi:
        db.session.add(Catatan(
            buku_id=int(buku_id), 
            bagian=bagian, 
            judul_bab=judul_bab, 
            isi=isi,
            urutan=urutan_val
        ))
        db.session.commit()
        catat_log("TAMBAH BAB", f"Menambahkan bab: {judul_bab}")
    return redirect(f'/buku/{buku_id}')

@app.route('/edit-catatan/<int:catatan_id>', methods=['POST'])
def edit_catatan(catatan_id):
    if not session.get('is_admin'): return "Akses Ditolak", 403
    catatan = Catatan.query.get_or_404(catatan_id)
    catatan.bagian = request.form.get('bagian', '').strip()
    catatan.judul_bab = request.form.get('judul_bab', '').strip()
    catatan.isi = request.form.get('isi', '').strip()
    
    urutan_val = request.form.get('urutan')
    if urutan_val:
        try:
            catatan.urutan = int(urutan_val)
        except ValueError:
            pass

    db.session.commit()
    catat_log("EDIT BAB", f"Memperbarui bab: {catatan.judul_bab}")
    return redirect(f'/buku/{catatan.buku_id}')

@app.route('/hapus-catatan/<int:catatan_id>/<int:buku_id>', methods=['POST'])
def hapus_catatan(catatan_id, buku_id):
    if not session.get('is_admin'): return "Akses Ditolak", 403
    catatan = Catatan.query.get(catatan_id)
    if catatan:
        db.session.delete(catatan)
        db.session.commit()
    return redirect(f'/buku/{buku_id}')

@app.route('/export-buku/<int:buku_id>')
def export_buku(buku_id):
    buku = Buku.query.get_or_404(buku_id)
    teks = f"JUDUL: {buku.judul}\n" + "\n".join([f"[{c.bagian}] {c.judul_bab}\n{c.isi}\n" for c in buku.catatan_list])
    return Response(teks, mimetype='text/plain', headers={'Content-Disposition': f'attachment;filename={buku.judul}.txt'})

@app.route('/export-buku-docx/<int:buku_id>')
def export_buku_docx(buku_id):
    buku = Buku.query.get_or_404(buku_id)
    doc = Document()
    doc.add_heading(buku.judul.upper(), level=0)
    if buku.subjudul: doc.add_paragraph(buku.subjudul)
    doc.add_page_break()

    for c in buku.catatan_list:
        doc.add_heading(c.judul_bab, level=2)
        for p in c.isi.split('\n'):
            if p.strip(): doc.add_paragraph(p.strip())

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    filename = re.sub(r'[^a-zA-Z0-9_.-]', '_', buku.judul)
    return send_file(buffer, as_attachment=True, download_name=f"{filename}.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/export-buku-epub/<int:buku_id>')
def export_buku_epub(buku_id):
    buku = Buku.query.get_or_404(buku_id)
    book = epub.EpubBook()
    book.set_identifier(f'buku-{buku.id}')
    book.set_title(buku.judul)
    book.set_language('id')
    book.add_author('Dede Suhendra')

    chapters = []
    for idx, c in enumerate(buku.catatan_list):
        ch = epub.EpubHtml(title=c.judul_bab, file_name=f'chap_{idx+1}.xhtml', lang='id')
        isi_html = f"<h2>{escape(c.judul_bab)}</h2>"
        for p in c.isi.split('\n'):
            if p.strip(): isi_html += f"<p>{escape(p.strip())}</p>"
        ch.content = isi_html
        book.add_item(ch)
        chapters.append(ch)

    book.toc = tuple(chapters)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ['nav'] + chapters

    buffer = BytesIO()
    epub.write_epub(buffer, book, {})
    buffer.seek(0)
    filename = re.sub(r'[^a-zA-Z0-9_.-]', '_', buku.judul)
    return send_file(buffer, as_attachment=True, download_name=f"{filename}.epub", mimetype='application/epub+zip')

# Prioritas Tinggi: Ekspor PDF Terproteksi Anti-Crash HTML
@app.route('/export-buku-pdf/<int:buku_id>')
def export_buku_pdf(buku_id):
    buku = Buku.query.get_or_404(buku_id)
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
    
    styles = getSampleStyleSheet()
    
    style_cover_title = ParagraphStyle('CoverTitle', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=24, leading=30, alignment=TA_CENTER)
    style_cover_sub = ParagraphStyle('CoverSub', parent=styles['Normal'], fontName='Helvetica', fontSize=13, leading=18, alignment=TA_CENTER)
    style_cover_quote = ParagraphStyle('CoverQuote', parent=styles['Normal'], fontName='Helvetica-Oblique', fontSize=10, leading=14, alignment=TA_CENTER)
    style_bagian = ParagraphStyle('BagianHeader', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=13, leading=17, spaceBefore=15, spaceAfter=6)
    style_bab_title = ParagraphStyle('BabTitle', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=11, leading=15, spaceBefore=10, spaceAfter=4)
    style_bab_body = ParagraphStyle('BabBody', parent=styles['Normal'], fontName='Helvetica', fontSize=10, leading=14, alignment=TA_JUSTIFY, spaceAfter=10)

    story = []
    story.append(Spacer(1, 120))
    story.append(Paragraph(f"<b>{escape(buku.judul.upper())}</b>", style_cover_title))
    story.append(Spacer(1, 15))
    if buku.subjudul:
        story.append(Paragraph(escape(buku.subjudul), style_cover_sub))
        story.append(Spacer(1, 20))
    story.append(Paragraph("Disusun oleh: Dede Suhendra", style_cover_sub))
    story.append(Spacer(1, 40))
    if buku.kutipan:
        story.append(Paragraph(f'"{escape(buku.kutipan)}"', style_cover_quote))
    story.append(PageBreak())

    last_bagian = None
    for c in buku.catatan_list:
        if c.bagian and c.bagian != last_bagian:
            last_bagian = c.bagian
            story.append(Paragraph(f"<b>{escape(last_bagian)}</b>", style_bagian))
        
        story.append(Paragraph(f"<b>{escape(c.judul_bab)}</b>", style_bab_title))
        isi_pdf_aman = bersihkan_teks_pdf(c.isi)
        story.append(Paragraph(isi_pdf_aman, style_bab_body))
        story.append(Spacer(1, 8))

    doc.build(story)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name=f"{re.sub(r'[^a-zA-Z0-9_.-]', '_', buku.judul)}.pdf", mimetype='application/pdf')

@app.route('/statistik')
def statistik():
    return render_template_string(HTML_STATISTIK, total_buku=Buku.query.count(), total_catatan=Catatan.query.count())

@app.route('/tong-sampah')
def tong_sampah(): return render_template_string(HTML_TONG_SAMPAH, sampah_list=TongSampah.query.all())

@app.route('/riwayat-log')
def riwayat_log():
    if not session.get('is_admin'): return redirect('/login')
    return render_template_string(HTML_LOG, logs=LogAktivitas.query.order_by(LogAktivitas.id.desc()).limit(50).all())

@app.route('/pulihkan/<int:item_id>')
def pulihkan(item_id):
    if not session.get('is_admin'): return "Akses Ditolak", 403
    item = TongSampah.query.get_or_404(item_id)
    db.session.delete(item)
    db.session.commit()
    return redirect('/tong-sampah')

@app.route('/hapus-permanen/<int:item_id>')
def hapus_permanen(item_id):
    if not session.get('is_admin'): return "Akses Ditolak", 403
    db.session.delete(TongSampah.query.get_or_404(item_id))
    db.session.commit()
    return redirect('/tong-sampah')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        if not cek_rate_limit_login(request.remote_addr):
            return "Terlalu banyak percobaan login. Coba lagi 1 menit kemudian.", 429
        if request.form.get('username') == ADMIN_USER and check_password_hash(ADMIN_PASS_HASH, request.form.get('password')):
            session['is_admin'] = True
            return redirect('/')
        return render_template_string(HTML_LOGIN, error="Salah password!")
    return render_template_string(HTML_LOGIN, error=None)

@app.route('/logout')
def logout():
    session.pop('is_admin', None)
    return redirect('/')

@app.route('/backup')
def backup_db():
    if not session.get('is_admin'): return "Akses Ditolak", 403
    data = {
        "tema": [{"id": t.id, "nama": t.nama} for t in Tema.query.all()],
        "buku": [{"id": b.id, "judul": b.judul, "subjudul": b.subjudul, "tema_id": b.tema_id, "kutipan": b.kutipan} for b in Buku.query.all()],
        "catatan": [{"id": c.id, "bagian": c.bagian, "judul_bab": c.judul_bab, "isi": c.isi, "buku_id": c.buku_id, "urutan": c.urutan} for c in Catatan.query.all()],
        "esai": [{"id": e.id, "judul": e.judul, "kategori": e.kategori, "isi": e.isi} for e in EsaiPenulis.query.all()]
    }
    return Response(json.dumps(data, indent=2), mimetype='application/json', headers={'Content-Disposition': 'attachment;filename=backup_karya.json'})

@app.route('/restore', methods=['POST'])
def restore_db():
    if not session.get('is_admin'):
        return "Akses Ditolak", 403
    file = request.files.get('file_json')
    if file:
        try:
            data = json.load(file)
            for t in data.get('tema', []):
                if not Tema.query.get(t['id']):
                    db.session.add(Tema(id=t['id'], nama=t['nama']))
            for b in data.get('buku', []):
                if not Buku.query.get(b['id']):
                    db.session.add(Buku(id=b['id'], judul=b['judul'], subjudul=b.get('subjudul'), tema_id=b['tema_id'], kutipan=b.get('kutipan')))
            for c in data.get('catatan', []):
                if not Catatan.query.get(c['id']):
                    db.session.add(Catatan(id=c['id'], bagian=c.get('bagian'), judul_bab=c['judul_bab'], isi=c['isi'], buku_id=c['buku_id'], urutan=c.get('urutan', 1)))
            for e in data.get('esai', []):
                if not EsaiPenulis.query.get(e['id']):
                    db.session.add(EsaiPenulis(id=e['id'], judul=e['judul'], kategori=e.get('kategori', 'Refleksi'), isi=e['isi']))
            db.session.commit()
            catat_log("RESTORE DATABASE", "Melakukan restore data dari file JSON")
            app.logger.info("Database berhasil di-restore dari file JSON.")
        except Exception as err:
            db.session.rollback()
            app.logger.error(f"Gagal restore database: {str(err)}")
    return redirect('/')

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)
